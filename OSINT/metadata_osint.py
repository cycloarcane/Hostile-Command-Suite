#!/usr/bin/env python3
"""
metadata_osint.py — FastMCP tool for extracting metadata from files and images

FastMCP tools
────────────
    extract_image_metadata(file_path, ...)
    extract_document_metadata(file_path, ...)
    extract_metadata_from_url(url, ...)
    search_files_by_metadata(directory, criteria, ...)

Returns
───────
    {
      "status": "success",
      "file_path": "...",
      "metadata": {...},
      "gps_coordinates": {...},
      "privacy_risks": [...]
    }

Dependencies
────────────
    pip install pillow exifread python-docx PyPDF2 python-magic-bin requests

Setup
─────
    Ensure ExifTool is installed for advanced metadata extraction:
    - Windows: Download from https://exiftool.org/
    - Linux: sudo apt install libimage-exiftool-perl
    - macOS: brew install exiftool
"""

import json
import os
import sys
import subprocess
import shutil
import tempfile
import logging
from typing import Any, Dict, List, Optional, Union
from urllib.parse import urlparse
import hashlib
import time

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    import requests
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import exifread
    EXIFREAD_AVAILABLE = True
except ImportError:
    EXIFREAD_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

from fastmcp import FastMCP

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("metadata_osint")

mcp = FastMCP("metadata")  # MCP route → /metadata
_CACHE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "cache")
os.makedirs(_CACHE_DIR, exist_ok=True)


def _check_dependencies() -> Dict[str, bool]:
    """Check which dependencies are available"""
    return {
        "pillow": PIL_AVAILABLE,
        "exifread": EXIFREAD_AVAILABLE,
        "pypdf2": PYPDF2_AVAILABLE,
        "docx": DOCX_AVAILABLE,
        "magic": MAGIC_AVAILABLE,
        "exiftool": shutil.which("exiftool") is not None
    }


def _get_file_type(file_path: str) -> str:
    """Determine file type using magic or file extension"""
    if MAGIC_AVAILABLE:
        try:
            mime = magic.Magic(mime=True)
            return mime.from_file(file_path)
        except Exception:
            pass
    
    # Fallback to extension
    _, ext = os.path.splitext(file_path.lower())
    ext_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.gif': 'image/gif',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
    }
    return ext_map.get(ext, 'application/octet-stream')


def _extract_gps_coordinates(exif_data: Dict) -> Optional[Dict[str, float]]:
    """Extract GPS coordinates from EXIF data"""
    try:
        gps_info = exif_data.get('GPSInfo')
        if not gps_info:
            return None
        
        def convert_to_degrees(value):
            """Convert GPS coordinates to degrees"""
            if isinstance(value, str):
                return float(value)
            
            d, m, s = value
            return d + (m / 60.0) + (s / 3600.0)
        
        gps_data = {}
        
        # Extract latitude
        if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
            lat = convert_to_degrees(gps_info['GPSLatitude'])
            if gps_info['GPSLatitudeRef'] != 'N':
                lat = -lat
            gps_data['latitude'] = lat
        
        # Extract longitude
        if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
            lon = convert_to_degrees(gps_info['GPSLongitude'])
            if gps_info['GPSLongitudeRef'] != 'E':
                lon = -lon
            gps_data['longitude'] = lon
        
        # Extract altitude
        if 'GPSAltitude' in gps_info:
            alt = float(gps_info['GPSAltitude'])
            if gps_info.get('GPSAltitudeRef', 0) == 1:
                alt = -alt
            gps_data['altitude'] = alt
        
        return gps_data if gps_data else None
        
    except Exception as e:
        logger.warning(f"Error extracting GPS data: {e}")
        return None


def _analyze_privacy_risks(metadata: Dict) -> List[str]:
    """Analyze metadata for privacy risks"""
    risks = []
    
    # Check for GPS coordinates
    if metadata.get('gps_coordinates'):
        risks.append("GPS coordinates found - reveals location where photo was taken")
    
    # Check for camera information
    if metadata.get('camera_make') or metadata.get('camera_model'):
        risks.append("Camera information found - can be used for device fingerprinting")
    
    # Check for software information
    if metadata.get('software'):
        risks.append("Software information found - reveals editing tools used")
    
    # Check for timestamps
    if metadata.get('datetime_original') or metadata.get('datetime'):
        risks.append("Timestamp information found - reveals when content was created")
    
    # Check for author/creator information
    if metadata.get('artist') or metadata.get('author') or metadata.get('creator'):
        risks.append("Author/creator information found - personal identification possible")
    
    # Check for comments
    if metadata.get('user_comment') or metadata.get('comment'):
        risks.append("Comments found - may contain sensitive information")
    
    return risks


@mcp.tool()
def check_metadata_tools() -> Dict[str, Union[str, Dict[str, bool]]]:
    """
    Check which metadata extraction tools and dependencies are available.
    
    Returns:
        A dictionary with tool availability status
    """
    deps = _check_dependencies()
    
    missing = []
    if not deps["pillow"]:
        missing.append("pillow: pip install pillow")
    if not deps["exifread"]:
        missing.append("exifread: pip install exifread")
    if not deps["pypdf2"]:
        missing.append("pypdf2: pip install PyPDF2")
    if not deps["docx"]:
        missing.append("docx: pip install python-docx")
    if not deps["magic"]:
        missing.append("magic: pip install python-magic-bin")
    if not deps["exiftool"]:
        missing.append("exiftool: Install from https://exiftool.org/")
    
    return {
        "status": "ok" if all(deps.values()) else "missing_dependencies",
        "dependencies": deps,
        "installation_instructions": missing
    }


@mcp.tool()
def extract_image_metadata(
    file_path: str,
    include_thumbnail: bool = False,
    privacy_check: bool = True
) -> Dict[str, Any]:
    """
    Extract metadata from image files.
    
    Args:
        file_path: Path to the image file
        include_thumbnail: Whether to extract thumbnail data
        privacy_check: Whether to analyze privacy risks
        
    Returns:
        Dict with extracted metadata and analysis
    """
    if not os.path.exists(file_path):
        return {
            "status": "error",
            "message": f"File not found: {file_path}"
        }
    
    result = {
        "status": "success",
        "file_path": file_path,
        "file_size": os.path.getsize(file_path),
        "file_type": _get_file_type(file_path),
        "metadata": {},
        "gps_coordinates": None,
        "privacy_risks": []
    }
    
    try:
        # Method 1: PIL/Pillow extraction
        if PIL_AVAILABLE:
            try:
                with Image.open(file_path) as img:
                    result["metadata"]["image_size"] = img.size
                    result["metadata"]["image_mode"] = img.mode
                    result["metadata"]["image_format"] = img.format
                    
                    # Extract EXIF data
                    exif_data = img._getexif()
                    if exif_data:
                        exif_dict = {}
                        for tag, value in exif_data.items():
                            tag_name = TAGS.get(tag, tag)
                            
                            # Handle GPS info specially
                            if tag_name == "GPSInfo":
                                gps_dict = {}
                                for gps_tag in value:
                                    sub_tag_name = GPSTAGS.get(gps_tag, gps_tag)
                                    gps_dict[sub_tag_name] = value[gps_tag]
                                exif_dict[tag_name] = gps_dict
                            else:
                                # Convert to string if it's not JSON serializable
                                try:
                                    json.dumps(value)
                                    exif_dict[tag_name] = value
                                except (TypeError, ValueError):
                                    exif_dict[tag_name] = str(value)
                        
                        result["metadata"]["exif"] = exif_dict
                        
                        # Extract GPS coordinates
                        result["gps_coordinates"] = _extract_gps_coordinates(exif_dict)
                        
            except Exception as e:
                logger.warning(f"PIL extraction failed: {e}")
        
        # Method 2: ExifRead extraction (more detailed)
        if EXIFREAD_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    tags = exifread.process_file(f, details=True)
                    
                    exif_read_data = {}
                    for tag, value in tags.items():
                        if tag not in ('JPEGThumbnail', 'TIFFThumbnail', 'Filename', 'EXIF MakerNote'):
                            try:
                                exif_read_data[tag] = str(value)
                            except:
                                continue
                    
                    if exif_read_data:
                        result["metadata"]["exifread"] = exif_read_data
                        
            except Exception as e:
                logger.warning(f"ExifRead extraction failed: {e}")
        
        # Method 3: ExifTool extraction (most comprehensive)
        if shutil.which("exiftool"):
            try:
                cmd = ["exiftool", "-json", "-n", file_path]
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if proc.returncode == 0:
                    exiftool_data = json.loads(proc.stdout)
                    if exiftool_data and isinstance(exiftool_data, list):
                        result["metadata"]["exiftool"] = exiftool_data[0]
                        
                        # Extract additional GPS info from ExifTool
                        tool_data = exiftool_data[0]
                        if not result["gps_coordinates"] and any(k.startswith("GPS") for k in tool_data.keys()):
                            gps_data = {}
                            if "GPSLatitude" in tool_data and "GPSLongitude" in tool_data:
                                gps_data["latitude"] = tool_data["GPSLatitude"]
                                gps_data["longitude"] = tool_data["GPSLongitude"]
                            if "GPSAltitude" in tool_data:
                                gps_data["altitude"] = tool_data["GPSAltitude"]
                            if gps_data:
                                result["gps_coordinates"] = gps_data
                                
            except Exception as e:
                logger.warning(f"ExifTool extraction failed: {e}")
        
        # Consolidate common metadata fields
        consolidated = {}
        all_metadata = result["metadata"]
        
        # Try to find common fields across different extraction methods
        for source_data in all_metadata.values():
            if isinstance(source_data, dict):
                # Camera information
                for key in ["Make", "Model", "Camera", "camera_make", "camera_model"]:
                    if key in source_data and not consolidated.get("camera_make"):
                        consolidated["camera_make"] = source_data[key]
                
                # DateTime
                for key in ["DateTime", "DateTimeOriginal", "CreateDate", "datetime", "datetime_original"]:
                    if key in source_data and not consolidated.get("datetime_original"):
                        consolidated["datetime_original"] = source_data[key]
                
                # Software
                for key in ["Software", "software"]:
                    if key in source_data and not consolidated.get("software"):
                        consolidated["software"] = source_data[key]
                
                # Artist/Author
                for key in ["Artist", "Author", "Creator", "artist", "author"]:
                    if key in source_data and not consolidated.get("artist"):
                        consolidated["artist"] = source_data[key]
        
        result["metadata"]["consolidated"] = consolidated
        
        # Privacy risk analysis
        if privacy_check:
            result["privacy_risks"] = _analyze_privacy_risks({
                **consolidated,
                "gps_coordinates": result["gps_coordinates"]
            })
        
        return result
        
    except Exception as e:
        return {
            "status": "error",
            "file_path": file_path,
            "message": f"Error extracting metadata: {str(e)}"
        }


@mcp.tool()
def extract_document_metadata(
    file_path: str,
    privacy_check: bool = True
) -> Dict[str, Any]:
    """
    Extract metadata from document files (PDF, Word, etc.).
    
    Args:
        file_path: Path to the document file
        privacy_check: Whether to analyze privacy risks
        
    Returns:
        Dict with extracted metadata and analysis
    """
    if not os.path.exists(file_path):
        return {
            "status": "error",
            "message": f"File not found: {file_path}"
        }
    
    result = {
        "status": "success",
        "file_path": file_path,
        "file_size": os.path.getsize(file_path),
        "file_type": _get_file_type(file_path),
        "metadata": {},
        "privacy_risks": []
    }
    
    file_type = _get_file_type(file_path)
    
    try:
        # PDF extraction
        if "pdf" in file_type.lower() and PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    
                    if reader.metadata:
                        pdf_metadata = {}
                        for key, value in reader.metadata.items():
                            # Remove the leading slash from PDF metadata keys
                            clean_key = key.lstrip('/')
                            pdf_metadata[clean_key] = str(value) if value else ""
                        
                        result["metadata"]["pdf"] = pdf_metadata
                        result["metadata"]["page_count"] = len(reader.pages)
                        
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
        
        # Word document extraction
        elif "wordprocessingml" in file_type.lower() and DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                
                docx_metadata = {}
                core_props = doc.core_properties
                
                # Extract core properties
                for prop in ['author', 'category', 'comments', 'content_status', 
                            'created', 'identifier', 'keywords', 'language', 
                            'last_modified_by', 'last_printed', 'modified', 
                            'revision', 'subject', 'title', 'version']:
                    value = getattr(core_props, prop, None)
                    if value:
                        docx_metadata[prop] = str(value)
                
                if docx_metadata:
                    result["metadata"]["docx"] = docx_metadata
                
                # Count paragraphs and tables
                result["metadata"]["paragraph_count"] = len(doc.paragraphs)
                result["metadata"]["table_count"] = len(doc.tables)
                
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
        
        # ExifTool extraction (works for many file types)
        if shutil.which("exiftool"):
            try:
                cmd = ["exiftool", "-json", file_path]
                proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if proc.returncode == 0:
                    exiftool_data = json.loads(proc.stdout)
                    if exiftool_data and isinstance(exiftool_data, list):
                        result["metadata"]["exiftool"] = exiftool_data[0]
                        
            except Exception as e:
                logger.warning(f"ExifTool extraction failed: {e}")
        
        # Consolidate metadata
        consolidated = {}
        all_metadata = result["metadata"]
        
        for source_data in all_metadata.values():
            if isinstance(source_data, dict):
                # Author information
                for key in ["Author", "Creator", "author", "creator"]:
                    if key in source_data and not consolidated.get("author"):
                        consolidated["author"] = source_data[key]
                
                # Title
                for key in ["Title", "title"]:
                    if key in source_data and not consolidated.get("title"):
                        consolidated["title"] = source_data[key]
                
                # Creation date
                for key in ["CreationDate", "CreateDate", "created", "Created"]:
                    if key in source_data and not consolidated.get("created"):
                        consolidated["created"] = source_data[key]
                
                # Modification date
                for key in ["ModifyDate", "modified", "Modified", "last_modified_by"]:
                    if key in source_data and not consolidated.get("modified"):
                        consolidated["modified"] = source_data[key]
                
                # Subject/Keywords
                for key in ["Subject", "Keywords", "subject", "keywords"]:
                    if key in source_data and not consolidated.get("subject"):
                        consolidated["subject"] = source_data[key]
                
                # Producer/Software
                for key in ["Producer", "Software", "Application", "producer", "software"]:
                    if key in source_data and not consolidated.get("producer"):
                        consolidated["producer"] = source_data[key]
        
        result["metadata"]["consolidated"] = consolidated
        
        # Privacy risk analysis
        if privacy_check:
            result["privacy_risks"] = _analyze_privacy_risks(consolidated)
        
        return result
        
    except Exception as e:
        return {
            "status": "error",
            "file_path": file_path,
            "message": f"Error extracting metadata: {str(e)}"
        }


@mcp.tool()
def extract_metadata_from_url(
    url: str,
    download_timeout: int = 30,
    max_file_size: int = 50 * 1024 * 1024,  # 50MB
    privacy_check: bool = True
) -> Dict[str, Any]:
    """
    Download a file from URL and extract its metadata.
    
    Args:
        url: URL of the file to download and analyze
        download_timeout: Timeout for download in seconds
        max_file_size: Maximum file size to download in bytes
        privacy_check: Whether to analyze privacy risks
        
    Returns:
        Dict with extracted metadata and analysis
    """
    try:
        # Validate URL
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            return {
                "status": "error",
                "message": "Invalid URL provided"
            }
        
        # Download file to temporary location
        response = requests.get(url, stream=True, timeout=download_timeout)
        response.raise_for_status()
        
        # Check content length
        content_length = response.headers.get('Content-Length')
        if content_length and int(content_length) > max_file_size:
            return {
                "status": "error",
                "message": f"File too large: {content_length} bytes (max: {max_file_size})"
            }
        
        # Determine file extension from URL or Content-Type
        content_type = response.headers.get('Content-Type', '')
        
        extension_map = {
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/msword': '.doc'
        }
        
        ext = extension_map.get(content_type, os.path.splitext(parsed_url.path)[1])
        if not ext:
            ext = '.bin'
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as temp_file:
            downloaded_size = 0
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    downloaded_size += len(chunk)
                    if downloaded_size > max_file_size:
                        os.unlink(temp_file.name)
                        return {
                            "status": "error",
                            "message": f"File too large: {downloaded_size} bytes (max: {max_file_size})"
                        }
                    temp_file.write(chunk)
            
            temp_file_path = temp_file.name
        
        try:
            # Extract metadata based on content type
            if content_type.startswith('image/'):
                result = extract_image_metadata(temp_file_path, privacy_check=privacy_check)
            else:
                result = extract_document_metadata(temp_file_path, privacy_check=privacy_check)
            
            # Add URL information
            result["source_url"] = url
            result["content_type"] = content_type
            result["downloaded_size"] = downloaded_size
            
            return result
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        
    except requests.exceptions.RequestException as e:
        return {
            "status": "error",
            "message": f"Download failed: {str(e)}"
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Error processing file: {str(e)}"
        }


@mcp.tool()
def search_files_by_metadata(
    directory: str,
    criteria: Dict[str, str],
    recursive: bool = True,
    file_extensions: Optional[List[str]] = None
) -> Dict[str, Any]:
    """
    Search for files in a directory based on metadata criteria.
    
    Args:
        directory: Directory to search in
        criteria: Metadata criteria to match (e.g., {"author": "John Doe"})
        recursive: Whether to search subdirectories
        file_extensions: List of file extensions to check (e.g., [".jpg", ".pdf"])
        
    Returns:
        Dict with matching files and their metadata
    """
    if not os.path.exists(directory):
        return {
            "status": "error",
            "message": f"Directory not found: {directory}"
        }
    
    if not file_extensions:
        file_extensions = [".jpg", ".jpeg", ".png", ".gif", ".pdf", ".docx", ".doc"]
    
    matching_files = []
    total_files_checked = 0
    
    try:
        # Walk through directory
        for root, dirs, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                _, ext = os.path.splitext(file.lower())
                
                if ext in [e.lower() for e in file_extensions]:
                    total_files_checked += 1
                    
                    try:
                        # Extract metadata
                        if ext in ['.jpg', '.jpeg', '.png', '.gif']:
                            metadata_result = extract_image_metadata(file_path, privacy_check=False)
                        else:
                            metadata_result = extract_document_metadata(file_path, privacy_check=False)
                        
                        if metadata_result.get("status") == "success":
                            consolidated = metadata_result.get("metadata", {}).get("consolidated", {})
                            
                            # Check if file matches criteria
                            matches = True
                            for key, value in criteria.items():
                                file_value = consolidated.get(key, "")
                                if value.lower() not in str(file_value).lower():
                                    matches = False
                                    break
                            
                            if matches:
                                matching_files.append({
                                    "file_path": file_path,
                                    "metadata": consolidated,
                                    "file_size": metadata_result.get("file_size", 0)
                                })
                    
                    except Exception as e:
                        logger.warning(f"Error processing {file_path}: {e}")
                        continue
            
            # If not recursive, only check the first level
            if not recursive:
                break
        
        return {
            "status": "success",
            "directory": directory,
            "criteria": criteria,
            "total_files_checked": total_files_checked,
            "matching_files": matching_files,
            "matches_found": len(matching_files)
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Error searching directory: {str(e)}"
        }


if __name__ == "__main__":
    mcp.run(transport="stdio")